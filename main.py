import json
from docx import Document
import pandas as pd


def extract_tables_from_docx(filename: str) -> list:
    """
    Liest die vorliegend .docx-Datei und erstellt einen Dataframe aus der vorliegenden Tabelle. Wir nehmen an, dass die
    Datei nur eine Tabelle enthält.
    """
    doc = Document(filename)

    # Wir nehmen an, dass nur eine Tabelle existiert und verwenden daher doc.tables[0].
    table = doc.tables[0]
    data = []

    for i, row in enumerate(table.rows):
        row_data = [cell.text.strip() for cell in row.cells]  # .strip() entfernt überflüssige Leerzeichen

        # Wenn es die erste Zeile ist, nehmen wir an, dass es sich um die Spaltennamen handelt.
        if i == 0:
            columns = ['Index'] + row_data  # Füge 'Index' für den numerischen Index hinzu.
        else:
            data.append([i] + row_data)  # Füge den numerischen Index hinzu, der bei 1 beginnt.

    df = pd.DataFrame(data, columns=columns)

    df['Index'] = df['Index'].astype(int)  # Stellen Sie sicher, dass der Index numerisch ist

    # Caste den Dataframe zu einem Dictionary (List of Dictionary's).
    return df.to_dict('records')


def serialize(d: dict) -> None:
    """
    Speichert die Liste als json im Arbeitsverzeichnis.
    """
    with open('table.json', 'w', encoding='utf-8') as f:
        json.dump(d, f, ensure_ascii=False)


def deserialize() -> json:
    """
    Lädt die Tabelle, welche als JSON vorliegt und gibt sie als Objekt zurück.
    """
    return json.load(open("table.json", 'r', encoding='utf-8'))


def create_csv_from_dict(d: dict, create_seperate_tables=False) -> None:
    """ df = pd.DataFrame(columns=["gv_kurland_ueberlieferungen.archiv_neu", "gv_kurland_urkunden.druckdatum",
                               "gv_kurland_urkunden.sortierdatum", "gv_kurland_urkunden.ausort",
                               "gv_kurland_urkunden.regest",
                               "gv_kurland_ueberlieferungen.material", "gv_kurland_ueberlieferungen.art",
                               "gv_kurland_ueberlieferungen.siegel",
                               "gv_kurland_urkunden.urkunde.id", "gv_kurland_urkunden.anmerk_druck"])
    """
    table1 = []
    table2 = []
    table3 = []

    for e in d:
        result = distinguish_case(e)

        if result[0] == 1:
            table1.append(result[1])
        elif result[0] == 2:
            table2.append(result[1])
        elif result[0] == 3:
            table3.append(result[1])
        else:
            continue

    pd.DataFrame(table1).to_csv("Fall1.csv", index=False)
    pd.DataFrame(table2).to_csv("Fall2.csv", index=False)
    pd.DataFrame(table3).to_csv("Fall3.csv", index=False)


def distinguish_case(ce: dict) -> None or 0:
    if ce["Vorlage bekannt?"] == "" or ce["Vorlage bekannt?"].lower() == "nein" or ce["Vorlage bekannt?"] == "?":
        return 1, handle_case_one(ce)
    elif ce["Vorlage bekannt?"].lower() == "ja" and ce["KGU"].startswith("Bauer"):
        return 2, handle_case_two(ce)
    elif ce["Vorlage bekannt?"].lower() == "ja" and ce["KGU"].startswith("erg") and int(ce["KGU"][3:]) < 1500:
        return 3, handle_case_three(ce)
    else:
        print(f"Problem mit \n {ce}")
        return 0, None


def handle_case_one(ce: dict) -> dict:
    datum_und_ort = get_datum(ce)
    KGU = get_KGU(ce)
    Inhalt_und_Abschrift = get_inhalt(ce)

    new_line = {
        'gv_kurland_ueberlieferungen.archiv_neu': f"{ce['Signatur']}{', ' if ce['Signatur'] else ''}{ce['Blatt']}",
        'gv_kurland_urkunden.druckdatum': datum_und_ort['Druckdatum'],
        'gv_kurland_urkunden.sortierdatum': datum_und_ort['Sortierdatum'],
        'gv_kurland_urkunden.ausort': datum_und_ort['Austellungsort'],
        'gv_kurland_urkunden.regest': f"{ce['Inhalt'] if Inhalt_und_Abschrift is None else Inhalt_und_Abschrift[0]}",
        'gv_kurland_ueberlieferungen.material': ce['Material'],
        'gv_kurland_ueberlieferungen.art': f"{ce['O/K'] if Inhalt_und_Abschrift is None else Inhalt_und_Abschrift[1]}",
        'gv_kurland_ueberlieferungen.siegel': ce['Siegel'],
        'gv_kurland_urkunden.urkunde.id': f"{ce['KGU'] if KGU is None else KGU}",
        'gv_kurland_urkunden.anmerk_druck': f"{KGU if KGU is not None else ''}"
    }

    return new_line


def get_datum(ce: dict) -> dict:
    result_set = {
        "Druckdatum": "",
        "Sortierdatum": "",
        "Austellungsort": ""
    }
    line_lst = ce['Datum'].split("\n")
    first_line = line_lst[0]
    if "," in first_line:
        result_set["Sortierdatum"] = first_line[:(ce['Datum'].find(", "))]
        result_set["Austellungsort"] = first_line[(ce['Datum'].find(", ")) + 2:]
        if len(line_lst) > 1:
            remainder = ', '.join(line_lst[1:])
            result_set["Druckdatum"] = f"{result_set['Sortierdatum']} ({remainder})"
        else:
            result_set["Druckdatum"] = result_set["Sortierdatum"]
    else:
        result_set["Sortierdatum"] = first_line
        if len(line_lst) > 1:
            remainder = ', '.join(line_lst[1:])
            result_set["Druckdatum"] = f"{result_set['Sortierdatum']} ({remainder})"
        else:
            result_set["Druckdatum"] = result_set["Sortierdatum"]

    return result_set


def get_KGU(ce: dict) -> str or None:
    if "\n" in ce['KGU']:
        return " ".join(ce['KGU'].split("\n"))
    else:
        return None


def get_inhalt(ce: dict) -> list or None:
    if "\n" in ce['Inhalt']:
        return ce['Inhalt'].split("\n")
    else:
        return None


def handle_case_two(ce: dict) -> dict:
    datum_und_ort = get_datum(ce)
    Inhalt_und_Abschrift = get_inhalt(ce)

    new_line = {
        'gv_kurland_ueberlieferungen.archiv_neu': f"{ce['Signatur']}{', ' if ce['Signatur'] else ''}{ce['Blatt']}",
        'gv_kurland_urkunden.druckdatum': datum_und_ort['Druckdatum'],
        'gv_kurland_ueberlieferungen.material': ce['Material'],
        'gv_kurland_ueberlieferungen.art': f"{ce['O/K'] if Inhalt_und_Abschrift is None else Inhalt_und_Abschrift[1]}",
        'gv_kurland_ueberlieferungen.siegel': ce['Siegel']
    }

    return new_line


def handle_case_three(ce: dict) -> dict:
    datum_und_ort = get_datum(ce)
    Inhalt_und_Abschrift = get_inhalt(ce)

    new_line = {
        'gv_kurland_ueberlieferungen.archiv_neu': f"{ce['Signatur']}{', ' if ce['Signatur'] else ''}{ce['Blatt']}",
        'gv_kurland_urkunden.druckdatum': datum_und_ort['Druckdatum'],
        'gv_kurland_ueberlieferungen.material': ce['Material'],
        'gv_kurland_ueberlieferungen.art': f"{ce['O/K'] if Inhalt_und_Abschrift is None else Inhalt_und_Abschrift[1]}"
    }

    return new_line


def analyse_case_distribution(d):
    case1, case2, case3, case4 = 0, 0, 0, 0

    for e in d:

        if e["Vorlage bekannt?"] == "" or e["Vorlage bekannt?"].lower() == "nein" or e["Vorlage bekannt?"] == "?":
            case1 += 1
            pass
        elif e["Vorlage bekannt?"].lower() == "ja" and e["KGU"].startswith("Bauer"):
            case2 += 1
            pass
        elif e["Vorlage bekannt?"].lower() == "ja" and e["KGU"].startswith("erg") and int(e["KGU"][3:]) < 1500:
            case3 += 1
            pass
        else:
            case4 += 1
            print(e)

    print(
        f"{case1} Einträge erfüllen die erste Bedingung, {case2} Einträge die zweite Bedingung, {case3} Eintreäge die dritte Bedingung und {case4} Einträge die vierte Bedingung.")


def preliminary_work(d, inplace=True):
    counter = 0
    for e in d:
        if e['Material'] == "":
            e['Material'] = "Papier"
            d[e['Index'] - 1] = e

        if e['Inhalt'].count('\n') > 1:
            index_list = find_all_token(e['Inhalt'])
            e['O/K'] = e['Inhalt'][index_list[-1] + 1:]
            e['Inhalt'] = e['Inhalt'][:index_list[-1]]
            e = clean_line(e)
            d[e['Index'] - 1] = e

        if e['KGU'] == "":
            e['KGU'] = f"erg{1500 + counter}"
            d[e['Index'] - 1] = e
            counter += 1


def find_all_token(main_string, sub_string='\n'):
    start = 0
    result_set = []
    while start < len(main_string):
        start = main_string.find(sub_string, start)
        if start == -1:
            break
        result_set.append(start)
        start += 1

    return result_set


def clean_line(e: dict) -> dict:
    if e['Inhalt'].endswith(" ") or e['Inhalt'].endswith(","):
        e['Inhalt'] = e['Inhalt'][:-1]
    else:
        return e
    return clean_line(e)


if __name__ == '__main__':
    '''analyse_case_distribution(d)

            1365 Einträge erfüllen die erste Bedingung, 54 Einträge die zweite Bedingung, 21 Eintreäge die dritte Bedingung und 6 Einträge die vierte Bedingung.

            Problematische Einträge:
            {'Index': 1282, 'Signatur': 'Riga, Hist. StaatsA., Best. 6999: Dokumentensamlung der Liv-, Kur- und Estländischen Güter, Fb. 29, Brieflade Gut Schlockenbeck (Šlokenbeka), Akte 4', 'Blatt': 'Bl. 1, 2', 'Datum': '1531 Feb. 5, Wenden\nahm sondage nach purificationis Marie dem vifftenn dach des mantes February', 'Inhalt': 'Wolter von Plettenberg belehnt Hildebrand Brockhusen mit Land bei Tuckum\nAbschrift 19. Jh.', 'Material': '', 'O/K': 'K', 'Siegel': '', 'KGU': 'erg0852', 'Vorlage bekannt?': 'ja (aber mit Bezug auf Akte 20!)'}
            {'Index': 1291, 'Signatur': 'Riga, Hist. StaatsA., Best. 6999: Dokumentensamlung der Liv-, Kur- und Estländischen Güter, Fb. 29, Brieflade Gut Schlockenbeck (Šlokenbeka), Akte 21', 'Blatt': 'Bl. 1v, 2, 2v', 'Datum': '1531 Feb. 5, Wenden\nam Sonntage nach Mariae Reinigung den fünften Tag des Monaths Februarii', 'Inhalt': 'Wolter von Plettenberg belehnt Hildebrand Brockhusen mit Land bei Tuckum\nAbschrift 18. Jh.', 'Material': '', 'O/K': 'K', 'Siegel': '', 'KGU': 'erg0852', 'Vorlage bekannt?': 'ja (aber mit Bezug auf Akte 20!)'}
            {'Index': 1313, 'Signatur': 'Riga, Hist. StaatsA., Best. 6999: Dokumentensamlung der Liv-, Kur- und Estländischen Güter, Fb. 31, Brieflade Gut Stenden (Stende), Akte 26', 'Blatt': 'Bl.1', 'Datum': '1540 Apr. 12, Wenden\nden mondages nha Misericordias Domini', 'Inhalt': 'Hermann von Brüggenei belehnt Philipp von der Brüggen mit Senten... \nbeglaubigte Abschrift 1644', 'Material': '', 'O/K': 'K', 'Siegel': '', 'KGU': 'erg0347', 'Vorlage bekannt?': 'ja (mit Apr. 11!)'}
            {'Index': 1355, 'Signatur': 'Riga, Hist. StaatsA., Best. 6999: Dokumentensamlung der Liv-, Kur- und Estländischen Güter, Fb. 31, Brieflade Gut Stenden (Stende), Akte 1182', 'Blatt': 'Bl. 3 (2)', 'Datum': '1391 Feb. 24, Durben\nferia sexta ante dominica, qua cantatur Oculi', 'Inhalt': 'Wennemar von Brüggenei belehnt Winricus von Durben... \nAbschrift 16. Jh.', 'Material': '', 'O/K': 'K', 'Siegel': '', 'KGU': 'Bauer129\n(mit Hinweis auf Akte 1183)', 'Vorlage bekannt?': 'ja?'}
            {'Index': 1356, 'Signatur': 'Riga, Hist. StaatsA., Best. 6999: Dokumentensamlung der Liv-, Kur- und Estländischen Güter, Fb. 31, Brieflade Gut Stenden (Stende), Akte 1182', 'Blatt': 'Bl. 3 (3)', 'Datum': '1467 Sept. 8, Goldingen\nan dem daghe der geboerth der konnickliken maget Marien', 'Inhalt': 'Johann von Mengede belehnt Hermann Blomberg...\nAbschrift 16. Jh.', 'Material': '', 'O/K': 'K', 'Siegel': '', 'KGU': 'Bauer258\n(mit Hinweis auf Akte 1118)', 'Vorlage bekannt?': 'ja?'}
            {'Index': 1409, 'Signatur': 'Riga, Hist. StaatsA., Best. 7363: Dokumente der Baltischen Geschichte, Handschriftensammlung, Fb. 3, Akte 75 („Copialbuch aus dem XIVten Jahrh. von 1242-1353“/Kurzemes bīskāpijas koparijis/Kopiarium des Bistums Kurland, 13.-14. Jh., Pergament = „Goldinger Kopialbuch“)', 'Blatt': 'Bl. 10v, 11', 'Datum': '1253 Apr.\nin dem aprille', 'Inhalt': 'Bischof Heinrich von Kurland … \nÜbersetzung, Abschrift um 1340', 'Material': '', 'O/K': 'K', 'Siegel': '', 'KGU': 'Bauer020', 'Vorlage bekannt?': 'ja?'}
    '''

    '''dict1 = extract_tables_from_docx("LST_KurGüUrk_Überführung_final_V2(1).docx")
    preliminary_work(dict)
    serialize(dict1)'''

    dict = deserialize()
    print(json.dumps(dict, indent=4, ensure_ascii=False))

    create_csv_from_dict(dict)

